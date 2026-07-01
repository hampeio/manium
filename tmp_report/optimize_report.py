from __future__ import annotations

from copy import deepcopy
from hashlib import sha256
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = Path(r"C:\Users\面\Desktop\桥梁工程课程设计_项目1与项目2综合报告.docx")
OUT = Path(r"C:\Users\面\Documents\manium\桥梁工程课程设计_项目1与项目2综合报告_排版优化版.docx")

doc = Document(SRC)
body = doc.element.body
original_children = list(body.iterchildren())
cover_hash = sha256(b"".join(child.xml.encode("utf-8") for child in original_children[:18])).hexdigest()
original_nonempty = [p.text for p in doc.paragraphs if p.text.strip()]
original_table_text = [[cell.text for row in table.rows for cell in row.cells] for table in doc.tables]

INK = "243240"
MUTED = "5E6B75"
LINE = "7F8C96"
HEADER_FILL = "ECEFF1"
CALLOUT_FILL = "F3F5F6"


def set_east_asia(run, name="宋体", size=9.5, bold=None, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rpr.rFonts.set(qn("w:eastAsia"), name)
    rpr.rFonts.set(qn("w:ascii"), "Times New Roman")
    rpr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_margins(cell, top=55, start=80, bottom=55, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    old = tbl_pr.find(qn("w:tblBorders"))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:val"), "nil")
        borders.append(node)
    tbl_pr.append(borders)


def three_line_table(table):
    remove_table_borders(table)
    table.autofit = False
    rows = table.rows
    if not rows:
        return
    for ri, row in enumerate(rows):
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = OxmlElement("w:cantSplit")
        tr_pr.append(cant_split)
        for cell in row.cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, HEADER_FILL if ri == 0 else "FFFFFF")
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = tc_pr.find(qn("w:tcBorders"))
            if borders is None:
                borders = OxmlElement("w:tcBorders")
                tc_pr.append(borders)
            for child in list(borders):
                borders.remove(child)
            # Standard three-line table: top and header separator, plus final bottom line.
            if ri == 0:
                for edge, sz in (("top", "10"), ("bottom", "6")):
                    line = OxmlElement(f"w:{edge}")
                    line.set(qn("w:val"), "single")
                    line.set(qn("w:sz"), sz)
                    line.set(qn("w:color"), LINE)
                    borders.append(line)
            if ri == len(rows) - 1:
                line = OxmlElement("w:bottom")
                line.set(qn("w:val"), "single")
                line.set(qn("w:sz"), "10")
                line.set(qn("w:color"), LINE)
                borders.append(line)
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent = None
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.paragraph_format.keep_together = True
                p.paragraph_format.keep_with_next = ri < len(rows) - 1
                for run in p.runs:
                    set_east_asia(run, "黑体" if ri == 0 else "宋体", 8.5, ri == 0, INK)


def set_table_geometry(table, widths_cm):
    widths_dxa = [int(value / 2.54 * 1440) for value in widths_cm]
    total = sum(widths_dxa)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for index, width in enumerate(widths_dxa):
        if index < len(grid_cols):
            grid_cols[index].set(qn("w:w"), str(width))
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Cm(widths_cm[index])
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def has_drawing(paragraph):
    return bool(paragraph._p.xpath(".//w:drawing"))


def resize_drawing(paragraph, width_cm):
    drawings = paragraph._p.xpath(".//w:drawing")
    if not drawings:
        return
    inline = drawings[0].xpath(".//wp:inline | .//wp:anchor")
    if not inline:
        return
    extents = drawings[0].xpath(".//wp:extent")
    aexts = drawings[0].xpath(".//a:xfrm/a:ext")
    if not extents:
        return
    old_cx = int(extents[0].get("cx"))
    old_cy = int(extents[0].get("cy"))
    new_cx = int(width_cm * 360000)
    new_cy = int(old_cy * new_cx / old_cx)
    for ext in extents:
        ext.set("cx", str(new_cx))
        ext.set("cy", str(new_cy))
    for ext in aexts:
        ext.set("cx", str(new_cx))
        ext.set("cy", str(new_cy))


def find_caption(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise ValueError(prefix)


def preceding_drawing(caption):
    node = caption._p.getprevious()
    while node is not None:
        if node.tag == qn("w:p"):
            from docx.text.paragraph import Paragraph
            p = Paragraph(node, caption._parent)
            if has_drawing(p):
                return p
        node = node.getprevious()
    raise ValueError(f"No image before {caption.text}")


def add_label_paragraph(cell, label, text, first=False):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(label)
    set_east_asia(r, "黑体", 8.8, True, INK)
    r = p.add_run(text)
    set_east_asia(r, "宋体", 8.8, False, INK)
    return p


def screenshot_two_column(caption_prefix, notes):
    caption = find_caption(caption_prefix)
    image_p = preceding_drawing(caption)
    resize_drawing(image_p, 7.0)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7.25)
    table.columns[1].width = Cm(7.25)
    table._tbl.getparent().remove(table._tbl)
    image_p._p.addprevious(table._tbl)
    remove_table_borders(table)
    row = table.rows[0]
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))
    left, right = row.cells
    for c in (left, right):
        set_cell_margins(c, 60, 80, 60, 80)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    # Remove the empty default paragraph, then move the existing image paragraph unchanged.
    empty = left.paragraphs[0]._p
    left._tc.remove(empty)
    left._tc.append(image_p._p)
    image_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = right.paragraphs[0]
    title.paragraph_format.first_line_indent = None
    title.paragraph_format.space_after = Pt(4)
    rr = title.add_run("界面说明")
    set_east_asia(rr, "黑体", 10, True, INK)
    for i, (label, text) in enumerate(notes):
        add_label_paragraph(right, label, text, first=False)
    # Encourage Word/LibreOffice to keep the explanation block with the caption.
    for c in row.cells:
        for p in c.paragraphs:
            p.paragraph_format.keep_together = True
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(3)
    caption.paragraph_format.keep_together = True


def add_border_callout_before(target_heading, title, paragraphs):
    heading = next(p for p in doc.paragraphs if p.text.strip().startswith(target_heading))
    inserted = []
    p = heading.insert_paragraph_before()
    inserted.append(p)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_east_asia(r, "黑体", 10, True, INK)
    for text in paragraphs:
        p = heading.insert_paragraph_before()
        inserted.append(p)
        p.paragraph_format.first_line_indent = Cm(0.7)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.2
        r = p.add_run(text)
        set_east_asia(r, "宋体", 9.5, False, INK)
    for p in inserted:
        p_pr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CALLOUT_FILL)
        p_pr.append(shd)
        p_bdr = OxmlElement("w:pBdr")
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "14")
        left.set(qn("w:color"), "8796A5")
        left.set(qn("w:space"), "6")
        p_bdr.append(left)
        p_pr.append(p_bdr)


def add_io_note_after(heading_prefix, note):
    heading = next(p for p in doc.paragraphs if p.text.strip().startswith(heading_prefix))
    node = heading._p.getnext()
    from docx.text.paragraph import Paragraph
    last = heading
    while node is not None and node.tag == qn("w:p"):
        p = Paragraph(node, heading._parent)
        if p.style.name.startswith("Heading") or has_drawing(p):
            break
        last = p
        node = node.getnext()
    new = OxmlElement("w:p")
    last._p.addnext(new)
    p = Paragraph(new, last._parent)
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), CALLOUT_FILL)
    p_pr.append(shd)
    r = p.add_run("输入—处理—输出：")
    set_east_asia(r, "黑体", 9, True, MUTED)
    r = p.add_run(note)
    set_east_asia(r, "宋体", 9, False, MUTED)


# 1) Compact only the report body. Cover, main title, chapter titles and existing text remain untouched.
start_compacting = False
for p in doc.paragraphs:
    if p.text.strip().startswith("第一章"):
        start_compacting = True
    if not start_compacting:
        continue
    if p.style.name == "Heading 1":
        # Explicitly preserve chapter-title formatting.
        continue
    if p.style.name == "Heading 2":
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
    else:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.28
        if p.text.strip().startswith(("图", "表")):
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.keep_together = True
            if p.text.strip().startswith("表"):
                p.paragraph_format.keep_with_next = True

# 2) Standardize the seven existing data tables as compact three-line tables.
for table in doc.tables:
    three_line_table(table)
table_width_patterns = [
    [3.0, 5.2, 7.3],
    [1.8, 4.5, 9.2],
    [2.2, 4.4, 8.9],
    [2.3, 3.5, 4.8, 4.9],
    [2.2, 4.0, 4.0, 5.3],
    [3.4, 8.0, 4.1],
    [2.2, 6.5, 6.8],
]
for table, widths in zip(doc.tables, table_width_patterns):
    set_table_geometry(table, widths)

# 3) Use left-screenshot/right-explanation layout only for software screenshots.
screenshot_two_column("图4-1", [
    ("主要功能：", "集中完成桥梁主题输入、分镜参数设置、生成状态查看和片段管理。"),
    ("关键区域：", "左侧为输入与质量参数，中部为视频预览，下部为分镜、代码和日志。"),
    ("操作逻辑：", "提交主题后按阶段生成；用户选择片段，可查看代码并触发局部重渲染。"),
    ("课程作用：", "把桥梁构造、传力路径等课程内容组织为可检查的动画成果。"),
])
screenshot_two_column("图4-3", [
    ("主要功能：", "通过可连接节点表达输入、规划、分镜、代码、渲染和输出依赖。"),
    ("关键区域：", "左侧节点库，中部画布，右侧参数与生成结果，上方提供验证和运行。"),
    ("操作逻辑：", "先连接兼容端口，再验证DAG，最后按拓扑顺序执行并保存节点状态。"),
    ("课程作用：", "固化桥梁工程课程成果的生成步骤，便于教师检查设计逻辑。"),
])
screenshot_two_column("图4-4", [
    ("主要功能：", "在视频时间轴上提交矩形、箭头或画笔批注，并生成局部修改任务。"),
    ("关键区域：", "左侧为分镜信息与批注操作，右侧为视频画面、任务文本和修订记录。"),
    ("操作逻辑：", "暂停视频、添加编号标签、提交反馈；系统按时间匹配shot并构建上下文。"),
    ("课程作用：", "用于复核桥梁构件关系、荷载箭头和讲解顺序是否准确清晰。"),
])
# The manually supplied first version contains a hidden "n + spaces" residue in
# this caption. The user explicitly requested this example error to be fixed.
caption_44 = find_caption("图4-4")
caption_44_original = caption_44.text
if caption_44.runs:
    caption_44.runs[0].text = "图4-4 项目2交互式视频审阅界面"
    for run in caption_44.runs[1:]:
        run.text = ""
    set_east_asia(caption_44.runs[0], "宋体", 10, True, INK)
else:
    run = caption_44.add_run("图4-4 项目2交互式视频审阅界面")
    set_east_asia(run, "宋体", 10, True, INK)

# 4) Keep non-screenshot figures single-column, but modestly reduce oversized graphics.
single_sizes = {
    "图1-1": 13.2, "图2-1": 13.0, "图2-2": 13.2, "图3-1": 13.2,
    "图3-2": 12.4, "图4-2": 13.2, "图5-1": 13.2, "图5-2": 13.2,
    "图5-3": 13.2, "图7-1": 13.0,
}
for prefix, width in single_sizes.items():
    caption = find_caption(prefix)
    resize_drawing(preceding_drawing(caption), width)
    caption.paragraph_format.space_before = Pt(1)
    caption.paragraph_format.space_after = Pt(3)

# 5) Add factual, compact explanations without rewriting any existing sentence.
add_border_callout_before("5.3 ", "系统内部提示词链（补充说明）", [
    "系统并非使用一条提示词直接生成完整视频，而是把内部提示词分为七类：SYSTEM_PROMPT约束中文表达、ManimCE语法和安全边界；PLAN_AND_CODE_PROMPT用于中短任务的一体化规划；GENERATION_STRATEGY_PROMPT负责长课程的阶段与批次划分。",
    "STORYBOARD_BATCH_PROMPT把教学目标细化为不重复的分镜，并要求每镜明确对象、位置、进入、保留、退出和变换；CODE_FROM_PLAN_PROMPT依据完整计划生成整课代码；SEGMENT_CODE_PROMPT仅生成当前片段，保证局部重渲染边界。",
    "REPAIR_PROMPT接收教学目标、当前代码和检查或渲染错误，只修复实现问题而不改变桥梁工程教学内容。七类模板均可通过prompt_overrides.json覆盖，Prompt Store在启动时加载并过滤旧版英文模板，使提示词可维护、可追踪。",
    "采用分阶段提示词的原因是把“桥梁知识是否正确”“分镜顺序是否合理”“代码是否可运行”分开校核，减少一次性生成造成的主题漂移、旧素材复用和错误难定位问题。",
])
add_io_note_after("5.1 ", "输入为桥梁主题、素材和生成参数；处理过程固定问题边界并记录阶段事件；输出为problem_frame、manifest和JSONL日志。")
add_io_note_after("5.3 ", "输入为分镜与视觉计划；处理包括提示词约束、代码清洗和语法检查；输出为可独立渲染的ManimCE Scene源码。")
add_io_note_after("5.6 ", "输入为视频时刻、全局要求和编号批注；处理包括shot匹配、代码区域提取和对象映射；输出为context JSON与Agent Task。")
add_io_note_after("7.2 ", "输入为测试用例与样例工程；处理覆盖结构、服务、渲染、交互和回滚；输出为通过状态、日志及可复核结论。")

# 6) Tighten header/footer footprint without changing their text.
for section in doc.sections:
    section.header_distance = Cm(0.9)
    section.footer_distance = Cm(0.9)
    for p in section.header.paragraphs + section.footer.paragraphs:
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        for run in p.runs:
            if run.font.size is None or run.font.size.pt > 8.5:
                run.font.size = Pt(8.5)

# Preservation checks.
new_children = list(doc.element.body.iterchildren())
new_cover_hash = sha256(b"".join(child.xml.encode("utf-8") for child in new_children[:18])).hexdigest()
assert cover_hash == new_cover_hash, "Cover XML changed"

all_paragraph_text = [p.text for p in doc.paragraphs]
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            all_paragraph_text.extend(p.text for p in cell.paragraphs)
for text in original_nonempty:
    if text == caption_44_original:
        continue
    assert text in all_paragraph_text, f"Original paragraph missing: {text[:80]}"

for old_table in original_table_text:
    flattened = [cell.text for table in doc.tables for row in table.rows for cell in row.cells]
    for text in old_table:
        assert text in flattened, f"Original table cell missing: {text[:80]}"

doc.core_properties.title = "基于Manim的桥梁工程教学动画生成与交互审阅系统（排版优化版）"
doc.save(OUT)
print(OUT)
print(f"preserved_paragraphs={len(original_nonempty)}; original_tables={len(original_table_text)}; cover_hash={cover_hash}")
